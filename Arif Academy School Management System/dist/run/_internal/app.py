from flask import Flask, render_template, request, jsonify
from os import getenv
import os
import numpy as np
from datetime import datetime
import mysql.connector
from mysql.connector import Error
import pandas as pd
import datetime
from mysql.connector import errorcode
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from mysql.connector import IntegrityError
from openpyxl import load_workbook
from openpyxl.styles import Font, Alignment
import webbrowser
from threading import Timer
import subprocess

app = Flask(__name__)

def create_connection():
    try:
        connection = mysql.connector.connect(
            host='localhost',
            user='admin',
            password='root',
            database='arif_academy'
        )
        
        if connection.is_connected():
            return connection
                
    except Error as e:
        print("Error while connecting to MySQL", e)

@app.route('/')
def indexstudent():
    return render_template('index.html')

def create_fee_voucher_document(student_data,today):
    doc = Document()
    
    # Define styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    os.makedirs('vouchers', exist_ok=True)
    directory_path=f'vouchers/{student_data['Class']}'
    os.makedirs(directory_path, exist_ok=True)
    output_path=f'{directory_path}/fee_voucher_{student_data['Reg_No'].replace('/','')}.docx'

    # Define a method to add the content for one voucher
    def add_voucher_content(cell):
        cell._element.clear_content()
        cell.paragraphs.clear()

        cell.add_paragraph("Arif Academy of Education").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.add_paragraph("Building Academic Excellence").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        table = cell.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = f'S/No: {student_data['Reg_No']}'
        hdr_cells[1].text = f'Date: {str(today)}'

        cell.add_paragraph(f"Name: {student_data['Name']}")
        cell.add_paragraph(f"Class: {student_data['Class']}")

        particulars={
            'Admission Fee': '0',
            'Monthly Fee': str(student_data['Monthly_fee']),
            'Arrears': str(int(student_data['Dues']-int(student_data['Monthly_fee']))),
            'Annual Fund': '0',
            'Fines': '0'
        }
        if student_data['admission_fee_status']=='Unpaid':
            particulars['Admission Fee']=student_data['Admission_Fees']
        if student_data['annual_fund_status']=='Unpaid':
            particulars['Annual Fee']=student_data['Annual_Fund']

        # Add particulars table
        particulars_table = cell.add_table(rows=len(particulars) + 2, cols=3)
        particulars_table.autofit = True
        hdr_cells = particulars_table.rows[0].cells
        hdr_cells[0].text = 'Particulars'
        hdr_cells[1].text = 'Amount'
        hdr_cells[2].text = ''
        
        for i, (key, value) in enumerate(particulars.items(), start=1):
            row_cells = particulars_table.rows[i].cells
            row_cells[0].text = key
            row_cells[1].text = 'Rs.'
            row_cells[2].text = value
        
        # Total row
        row_cells = particulars_table.rows[-1].cells
        row_cells[0].text = 'Total'
        row_cells[1].text = 'Rs.'
        total_amount = sum(int(value) for value in particulars.values())
        row_cells[2].text = str(total_amount)

        cell.add_paragraph("Rs.(In words): _______________________")
        cell.add_paragraph("Sign. Class Teacher")

    # Create a table with two columns for two parallel copies
    outer_table = doc.add_table(rows=1, cols=2)
    outer_table.autofit = False

    # Add content for both columns
    for cell in outer_table.rows[0].cells:
        add_voucher_content(cell)
    
    # Save the filled document
    doc.save(output_path)
    return True


@app.route('/monthly_fee/api', methods=['POST'])
def fee_voucher_gen():
    connection=create_connection()
    cursor=connection.cursor()
    today=datetime.date.today()
    month=int(today.month)
    month=7
    try:
        cursor.execute(f'SELECT s.Registration_no, s.Name, s.Month_Fees, f.Dues, f.advance_payment, f.amount_paid, f.Status, f.advance_payments_last,f.advance_months, s.Class, s.Admission_Fees, s.Annual_Fund, s.admission_fee_status, s.annual_fund_status FROM students s LEFT JOIN (SELECT * FROM fee WHERE month={month}) f ON s.Registration_No=f.Registration_No WHERE f.Status="Unpaid";')
        data=cursor.fetchall()
        data_df=pd.DataFrame({"Reg_No":[str(i[0]) for i in data], "Name":[str(i[1]) for i in data], 
                        "Monthly_fee":[str(i[2]) for i in data],"Dues":[str(i[3]) for i in data],
                        "advance_payment":[str(i[4]) for i in data],"amount_paid":[str(i[5]) for i in data],
                        'status':[str(i[6]) for i in data],'advance_payments_last':[str(i[7]) for i in data],
                        'advance_months':[str(i[8]) for i in data],'Class':[str(i[9]) for i in data],
                        'Admission_Fees':[str(i[10]) for i in data],'Annual_Fund':[str(i[11]) for i in data],
                        'admission_fee_status':[str(i[12]) for i in data],'annual_fund_status':[str(i[13]) for i in data]})
        data_df['advance_payment']=data_df['advance_payment'].astype('int64')
        data_df['amount_paid']=data_df['amount_paid'].astype('int64')
        data_df['Dues']=data_df['Dues'].astype('int64')
        data_df['advance_payments_last']=data_df['advance_payments_last'].astype('int64')
        data_df['advance_months']=data_df['advance_months'].astype('int64')
        data_df['Monthly_fee']=data_df['Monthly_fee'].astype('int64')
        data_df['done']=data_df.apply(lambda row: create_fee_voucher_document(row, today),axis=1)
    except Exception as e:
        if 'connection' in locals() and connection.is_connected():
            cursor.close()
            connection.close()
        return jsonify({"result": str(e)+". Data not updated..."}), 500
    return jsonify({"result": "Vouchers Generated!!!"})

def teacher_update(data, cursor):
    if data['NAME'] is not None and data['NAME']!="":
        cursor.execute(f'UPDATE teachers SET Name = "{data['NAME']}" WHERE CNIC="{data['CNIC']}";')
    if data['SAL'] is not None and data['SAL']!="":
        cursor.execute(f'UPDATE teachers SET Salary = {data['SAL']} WHERE CNIC="{data['CNIC']}";')
    if data['CONTACT'] is not None and data['CONTACT']!="":
        cursor.execute(f'UPDATE teachers SET CONTACT = "{data['CONTACT']}" WHERE CNIC="{data['CNIC']}";')
    if data['DOJ'] is not None and data['DOJ']!="":
        cursor.execute(f'UPDATE teachers SET Date_OF_Joining = DATE "{data['DOJ']}" WHERE CNIC="{data['CNIC']}";')

@app.route('/teacher_data_update_delete/api', methods=['POST'])
def add_update_teacher():
    connection=create_connection()
    cursor=connection.cursor()

    data = request.get_json()
    data=dict(data)
    today=datetime.date.today()
    if int(data['INDEX'])==0:
        try:
            cursor.execute(f'INSERT INTO teachers (CNIC) VALUES("{data['CNIC']}");')
            teacher_update(data, cursor)
        except IntegrityError as e:
            if e.errno == 1062:  # Error code for duplicate entry
                if 'connection' in locals() and connection.is_connected():
                    cursor.close()
                    connection.close()
                return jsonify({"result": "CNIC you entered already exists"}), 400
            else:
                if 'connection' in locals() and connection.is_connected():
                    cursor.close()
                    connection.close()
                return jsonify({"result": str(e)}), 500
        except Exception as e:
            return jsonify({"result": str(e)}), 500
    
    elif int(data['INDEX'])==1:
        try:
            if data['newcnic'] is not None and data['newcnic']!='':
                cursor.execute(f'INSERT INTO teachers SELECT "{data['newcnic']}", Name, Salary, CONTACT, Date_OF_Joining FROM teachers WHERE CNIC = "{data['CNIC']}";')
                cursor.execute(f'UPDATE payment SET CNIC = "{data['newcnic']}" WHERE CNIC="{data['CNIC']}";')
                cursor.execute(f'DELETE FROM teachers WHERE CNIC="{data['CNIC']}";')
                data['CNIC']=data['newcnic']
            teacher_update(data, cursor)
        except Exception as e:
            return jsonify({"result": str(e)+". Data not updated..."}), 500
    elif int(data['INDEX'])==2:
        try:
            cursor.execute(f'DELETE FROM teachers WHERE CNIC="{data['CNIC']}";')
        except Exception as e:
            if 'connection' in locals() and connection.is_connected():
                cursor.close()
                connection.close()
            return jsonify({"result": str(e)+". Data not updated..."}), 500

    connection.commit()
    if 'connection' in locals() and connection.is_connected():
        cursor.close()
        connection.close()
    return jsonify({"result": "Data Updated Successfully!!!"})

def get_close(temp):
    if len(temp)>0:
        open=int(temp.iloc[-1,5])
    else:
        open=0
    return open

@app.route('/fee_payment/api', methods=['POST'])
def fee_payment():
    connection=create_connection()
    cursor=connection.cursor()

    data = request.get_json()
    data=dict(data)
    current_month = int(datetime.date.today().month)
    try:
        cursor.execute(f'SELECT * FROM fee WHERE month={current_month};')
        regs=cursor.fetchall()
        regs=[i for i in regs if str(i[4])==str(data['REG'])]
        if len(regs)>0:
            id=regs[0][0]
            dues=int(regs[0][2])
            reg=regs[0][4]
            cursor.execute(f'SELECT Month_Fees FROM students WHERE Registration_No="{str(reg)}"')
            monthlyfee=int(cursor.fetchall()[0][0])
            dues=dues-int(data["Total"])
            if dues<0:
                advp=int(data['advanced_payment'])
                advm=str(data['Advanced_month'])
                cursor.execute(f'UPDATE fee SET Dues=0, Status="Paid", advance={advp}, advance_month="{advm}" WHERE ID="{id}";')
            elif dues>0:
                cursor.execute(f'UPDATE fee SET Dues={dues} WHERE ID="{id}";')
            else:
                cursor.execute(f'UPDATE fee SET Dues=0, Status="Paid" WHERE ID="{id}";')
            today = datetime.date.today()
            if int(today.month)>6:
                filename=f"static/xlscsv/accounts_{today.year}_{int(today.year)+1}.csv"
                if data['month_spe'] is not None and data['month_spe']!='' and data['month_spe']!='0':
                    monthys=data['month_spe'].split(',')
                    inp=''
                    for i in monthys:
                        if int(data['month_spe'])>6:
                            val=data['month_spe']+'-'+str(today.year)
                        else:
                            val=data['month_spe']+'-'+str(today.year+1)
                        inp=','+str(val)
                    inp=inp[1:]
                    cursor.execute(f'UPDATE fee SET advance_specific_month="{inp}" WHERE ID="{id}";')       
            else:
                filename=f"static/xlscsv/accounts_{today.year-1}_{int(today.year)}.csv"
                if data['month_spe'] is not None and data['month_spe']!='' and data['month_spe']!='0':
                    monthys=data['month_spe'].split(',')
                    inp=''
                    data['Total']=int(data['Total']) - monthlyfee*len(monthys)
                    for i in monthys:
                        if int(data['month_spe'])>6:
                            val=data['month_spe']+'-'+str(today.year)
                        else:
                            val=data['month_spe']+'-'+str(today.year+1)
                        inp=','+str(val)
                    inp=inp[1:]
                    cursor.execute(f'UPDATE fee SET advance_specific_month="{inp}" WHERE ID="{id}";')

            connection.commit()
            if 'connection' in locals() and connection.is_connected():
                cursor.close()
                connection.close()
            df=pd.read_csv(filename)
            if data['Cash_Bank']=='Bank':
                cb='1010: Current (Bank Account)'
            else:
                cb='Cash'
            open=get_close(df[df['Head']==cb])
            open2=get_close(df[df['Head']=='1220: Fee Receivable'])
            close=int(data['Total'])+open
            close2=open2-int(data['Total'])
            new={"Date":[str(today),str(today)], "Head":[cb,'1220: Fee Receivable'],"Debit":[data['Total'],None],"Credit":[None,data['Total']],'Opening Bal':[open, open2],'Closing Bal':[close,close2]}
            df=pd.concat([df,pd.DataFrame(new)],axis=0)
            if data['month_spe'] is not None and data['month_spe']!='' and data['month_spe']!='0':
                open=get_close(df[df['Head']==cb])
                open2=get_close(df[df['Head']=='2200: Advanced Fee Received'])
                close=int(data['Total'])+monthlyfee*len(monthys)
                close2=open2+monthlyfee*len(monthys)
                new={"Date":[str(today),str(today)], "Head":[cb,'2200: Advanced Fee Received'],"Debit":[monthlyfee*len(monthys),None],"Credit":[None,monthlyfee*len(monthys)],'Opening Bal':[open, open2],'Closing Bal':[close,close2]}
            df.to_csv(filename, index=False)
        else:
            if 'connection' in locals() and connection.is_connected():
                cursor.close()
                connection.close()
            return jsonify({"result": "Student Registration Number not found"})
    except Exception as e:
        if 'connection' in locals() and connection.is_connected():
            cursor.close()
            connection.close()
        return jsonify({"result": str(e)+". Data not updated..."}), 500
    return jsonify({"result": f"Fee added for {data['NAME']}!!!"})

def update_students(data, cursor):
    if data['NAME'] is not None and data['NAME']!='':
        cursor.execute(f'UPDATE students SET Name = "{data['NAME']}" WHERE Registration_No="{data['REG']}";')
    if data['FNAME'] is not None and data['FNAME']!='':
        cursor.execute(f'UPDATE students SET FATHER_NAME = "{data['FNAME']}" WHERE Registration_No="{data['REG']}";')
    if data['GRNO'] is not None and data['GRNO']!='':
        cursor.execute(f'UPDATE students SET Gr_no = "{data['GRNO']}" WHERE Registration_No="{data['REG']}";')
    if data['DOB'] is not None and data['DOB']!='':
        cursor.execute(f'UPDATE students SET DOB = DATE "{data['DOB']}" WHERE Registration_No="{data['REG']}";')
    if data['CONTACT'] is not None and data['CONTACT']!='':
        cursor.execute(f'UPDATE students SET Contact = "{data['CONTACT']}" WHERE Registration_No="{data['REG']}";')
    if data['alt'] is not None and data['alt']!='':
        cursor.execute(f'UPDATE students SET Alternate_Contact = "{data['alt']}" WHERE Registration_No="{data['REG']}";')
    if data['admission'] is not None and data['admission']!='':
        cursor.execute(f'UPDATE students SET Admission_Fees = {data['admission']} WHERE Registration_No="{data['REG']}";')
    if data['annual'] is not None and data['annual']!='':
        cursor.execute(f'UPDATE students SET Annual_Fund = {data['annual']} WHERE Registration_No="{data['REG']}";')
    if data['B_FORM'] is not None and data['B_FORM']!='':
        cursor.execute(f'UPDATE students SET B_Form = "{data['BFORM']}" WHERE Registration_No="{data['REG']}";')
    if data['Fee'] is not None and data['Fee']!='':
        cursor.execute(f'UPDATE students SET Month_Fees = {data['Month_Fees']} WHERE Registration_No="{data['REG']}";')
    if data['CLASS'] is not None and data['CLASS']!='':
        cursor.execute(f'UPDATE students SET Class = "{data['CLASS']}" WHERE Registration_No="{data['REG']}";')
    if data['Annual_status'] is not None and data['Annual_status']!='':
        cursor.execute(f'UPDATE students SET annual_fund_status = "{data['Annual_status']}" WHERE Registration_No="{data['REG']}";')
    if data['Admission_status'] is not None and data['Admission_status']!='':
        cursor.execute(f'UPDATE students SET admission_fee_status = "{data['Admission_status']}" WHERE Registration_No="{data['REG']}";')

@app.route('/student_data_update_delete/api', methods=['POST'])
def add_student():
    connection=create_connection()
    cursor=connection.cursor()

    data = request.get_json()
    data=dict(data)
    print(data)
    try:
        if int(data['INDEX'])==0:
            cursor.execute(f'INSERT INTO students (Registration_No) VALUES ("{data['REG']}");')
            
            data['Annual_status']='Unpaid'
            data['Admission_status']='Unpaid'
            update_students(data, cursor)
            today = datetime.date.today()
            if int(today.month)>6:
                filename=f"static/xlscsv/accounts_{today.year}_{int(today.year)+1}.csv"
            else:
                filename=f"static/xlscsv/accounts_{today.year-1}_{int(today.year)}.csv"
            df=pd.read_csv(filename)
            open=get_close(df[df['Head']=='1220: Fee Receivable'])
            open2=get_close(df[df['Head']=='4020: Annual Fund Income'])
            close=int(data['annual'])+open
            open3=close
            open4=get_close(df[df['Head']=='4030: Admission Fee'])
            close2=open2+int(data['annual'])
            close3=open3+int(data['admission'])
            close4=open4+int(data['admission'])
            new={"Date":[str(today),str(today),str(today),str(today)], "Head":['1220: Fee Receivable','4020: Annual Fund Income','1220: Fee Receivable','4030: Admission Fee'],"Debit":[data['annual'],None,data['admission'],None],"Credit":[None,data['annual'],None,data['admission']],'Opening Bal':[open,open2,open3,open4],'Closing Bal':[close,close2,close3,close4]}
            df=pd.concat([df,pd.DataFrame(new)],axis=0)
            df.to_csv(filename, index=False)
            
        elif int(data['INDEX'])==1:
            update_students(data, cursor)
        elif int(data['INDEX'])==2:
            print("Yes")
            cursor.execute(f'DELETE FROM fee WHERE Registration_No="{data['REG']}";')
            cursor.execute(f'DELETE FROM students WHERE Registration_No="{data['REG']}";')

        ##############################################################################################################################################################################################################################################

        connection.commit()
        if 'connection' in locals() and connection.is_connected():
            cursor.close()
            connection.close()
    except Exception as e:
        if 'connection' in locals() and connection.is_connected():
            cursor.close()
            connection.close()
        return jsonify({"result": str(e)+". Data not updated..."}), 500
    return jsonify({"result": f"Data for {data['NAME']} Updated Successfully!!!"})

@app.route('/salary_payment/api', methods=['POST'])
def salary_payment():
    connection=create_connection()
    cursor=connection.cursor()
    data = request.get_json()
    if data['ADDITIONAL'] is None:
        data['ADDITIONAL']=0
    print(data)
    print("hello")
    cnic=data['CNIC']
    try:
        cursor.execute(f'SELECT * FROM teachers where CNIC="{cnic}";')
        cd=cursor.fetchall()
        cd=cd[0][0]
    except Exception as e:
        if 'connection' in locals() and connection.is_connected():
            cursor.close()
            connection.close()
        return jsonify({"result": str(e)+". CNIC not Found"})
    today=datetime.date.today()
    pid=(cnic+str(today)).replace('-','')
    try:
        cursor.execute(f'SELECT Salary FROM teachers WHERE CNIC="{cnic}"')
        basic=cursor.fetchall()
        basic=int(basic[0][0])
        total_payment=0
        if int(today.month)>6:
            filename=f"static/xlscsv/accounts_{today.year}_{int(today.year)+1}.csv"
        else:
            filename=f"static/xlscsv/accounts_{today.year-1}_{int(today.year)}.csv"
        df=pd.read_csv(filename)
        if data['MEDICAL'] is not None and data['MEDICAL']!='0':
            total_payment+=int(data['MEDICAL'])
            open=get_close(df[df['Head']=='5030: Medical Allowance'])
            open2=get_close(df[df['Head']=="2115: Accrued Employee Benefits"])
            close=int(data['MEDICAL'])+open
            close2=open2+int(data['MEDICAL'])
            open3=close2
            open4=get_close(df[df['Head']=="1010: Current (Bank Account)"])
            close3=open3-int(data['MEDICAL'])
            close4=open4-int(data['MEDICAL'])
            new={"Date":[today,today,today,today], "Head":['5030: Medical Allowance',"2115: Accrued Employee Benefits","2115: Accrued Employee Benefits","1010: Current (Bank Account)"],"Debit":[int(data['MEDICAL']),None,int(data['MEDICAL']),None],"Credit":[None,int(data['MEDICAL']),None,int(data['MEDICAL'])],'Opening Bal':[open,open2,open3,open4],'Closing Bal':[close,close2,close3,close4]}
            df=pd.concat([df,pd.DataFrame(new)],axis=0)
        if data['Eidi'] is not None and data['Eidi']!='0':
            total_payment+=int(data['Eidi'])
            open=get_close(df[df['Head']=='5040: Eidi / Special Bonus'])
            open2=get_close(df[df['Head']=="2115: Accrued Employee Benefits"])
            close=open+int(data['Eidi'])
            close2=open2+int(data['Eidi'])
            open3=close2
            open4=get_close(df[df['Head']=="1010: Current (Bank Account)"])
            close3=open3-int(data['Eidi'])
            close4=open4-int(data['Eidi'])
            new={"Date":[today,today,today,today], "Head":['5040: Eidi / Special Bonus',"2115: Accrued Employee Benefits","2115: Accrued Employee Benefits","1010: Current (Bank Account)"],"Debit":[int(data['Eidi']),None,int(data['Eidi']),None],"Credit":[None,int(data['Eidi']),None,int(data['Eidi'])],'Opening Bal':[open,open2,open3,open4],'Closing Bal':[close,close2,close3,close4]}
            df=pd.concat([df,pd.DataFrame(new)],axis=0)
        if data['ADVANCED'] is not None and data['ADVANCED']!='0':
            total_payment+=int(data['ADVANCED'])
            if int(data['ADVANCED'])<0:
                open=get_close(df[df['Head']=='1210: Accounts Receivable'])
                open2=get_close(df[df['Head']=="1010: Current (Bank Account)"])
                close=open+int(data['ADVANCED'])
                close2=open2-int(data['ADVANCED'])
                new={"Date":[today,today], "Head":['1210: Accounts Receivable',"1010: Current (Bank Account)"],"Debit":[int(data['ADVANCED']),None],"Credit":[None,int(data['ADVANCED'])],'Opening Bal':[open,open2],'Closing Bal':[close,close2]}
            elif int(data['ADVANCED'])>0:
                open=get_close(df[df['Head']=="1010: Current (Bank Account)"])
                open2=get_close(df[df['Head']=='1210: Accounts Receivable'])
                close=open+int(data['ADVANCED'])
                close2=open2-int(data['ADVANCED'])
                new={"Date":[today,today], "Head":["1010: Current (Bank Account)",'1210: Accounts Receivable'],"Debit":[int(data['ADVANCED']),None],"Credit":[None,int(data['ADVANCED'])],'Opening Bal':[open,open2],'Closing Bal':[close,close2]}
            df=pd.concat([df,pd.DataFrame(new)],axis=0)
        if data['ADDITIONAL'] is not None and data['ADDITIONAL']!='0':
            total_payment=int(data['ADDITIONAL'])
            basic+=int(data['ADDITIONAL'])
        if data['Att_allow'] is not None and data['Att_allow']!='0':
            total_payment=int(data['Att_allow'])
            basic+=int(data['Att_allow'])
        if int(data['ADVANCED'])<0:
            deductions=int(data['ADVANCED'])*(-1)
        else:
            deductions=0
        print(f'("{cnic}", "{pid}",{data['attendace']},{total_payment},{data['MEDICAL']},{int(today.month)},{deductions},{data['ADVANCED']},{data['Att_allow']},{data['ADDITIONAL']});')
        cursor.execute(f'INSERT INTO payment (CNIC, Payment_ID, Attendence, Total_Payment, Medical_Allowance, monthh, deductions, advance, att_allowance, add_allowance) VALUES("{cnic}", "{pid}",{data['attendace']},{total_payment},{data['MEDICAL']},{int(today.month)},{deductions},{data['ADVANCED']},{data['Att_allow']},{data['ADDITIONAL']});')
        connection.commit()
        if 'connection' in locals() and connection.is_connected():
            cursor.close()
            connection.close()
        open=get_close(df[df['Head']=='5000: Salaries and Wages'])
        open2=get_close(df[df['Head']=="2100: Accrued Salaries"])
        close=open+basic
        close2=open2+basic
        open3=close2
        open4=get_close(df[df['Head']=="1010: Current (Bank Account)"])
        close3=open3-basic
        close4=open4-basic
        new={"Date":[today,today,today,today], "Head":['5000: Salaries and Wages',"2100: Accrued Salaries","2100: Accrued Salaries","1010: Current (Bank Account)"],"Debit":[basic,None,basic,None],"Credit":[None,basic,None,basic],'Opening Bal':[open,open2,open3,open4],'Closing Bal':[close,close2,close3,close4]}
        df=pd.concat([df,pd.DataFrame(new)],axis=0)
        df.to_csv(filename, index=False)
    except Exception as e:
        if 'connection' in locals() and connection.is_connected():
            cursor.close()
            connection.close()
        return jsonify({"result": str(e)+". Data not updated..."}), 500    

    return jsonify({"result": "Payment Added Successfully!!!"})

@app.route('/general_voucher/api', methods=['POST'])
def general_voucher():
    connection = create_connection()
    cursor = connection.cursor()

    data = request.get_json()
    today = datetime.date.today()
    if int(today.month) > 6:
        filename = f"static/xlscsv/accounts_{today.year}_{int(today.year)+1}.csv"
    else:
        filename = f"static/xlscsv/accounts_{today.year-1}_{int(today.year)}.csv"
    df = pd.read_csv(filename)
    data['amount']=int(data['amount'])
    if (data['category']=='Assets'):
        if 'Contra' in data['head']:
            open=get_close(df[df['Head']==data['head']])
            close=open+int(data['amount'])
            if data['amount']>0:
                close=open+int(data['amount'])
                new={"Date":[str(today)], "Head":[data['head']],"Debit":[None],"Credit":[data['amount']],'Opening Bal':[open],'Closing Bal':[close]}
            else:
                close=open+int(data['amount'])
                new={"Date":[str(today)], "Head":[data['head']],"Debit":[data['amount']*(-1)],"Credit":[None],'Opening Bal':[open],'Closing Bal':[close]}
        else:
            open=get_close(df[df['Head']==data['head']])
            if data['amount']>0:
                close=open+int(data['amount'])
                new={"Date":[str(today)], "Head":[data['head']],"Debit":[data['amount']],"Credit":[None],'Opening Bal':[open],'Closing Bal':[close]}
            else:
                close=open+int(data['amount'])
                new={"Date":[str(today)], "Head":[data['head']],"Debit":[None],"Credit":[data['amount']*(-1)],'Opening Bal':[open],'Closing Bal':[close]}
    elif(data['category']=='Equity'):
        if 'Contra' in data['head']:
            open=get_close(df[df['Head']==data['head']])
            if data['amount']>0:
                close=open+int(data['amount'])
                new={"Date":[str(today)], "Head":[data['head']],"Debit":[data['amount']],"Credit":[None],'Opening Bal':[open],'Closing Bal':[close]}
            else:
                close=open+int(data['amount'])
                new={"Date":[str(today)], "Head":[data['head']],"Debit":[None],"Credit":[data['amount']*(-1)],'Opening Bal':[open],'Closing Bal':[close]}
        else:
            open=get_close(df[df['Head']==data['head']])
            if data['amount']>0:
                close=open+int(data['amount'])
                new={"Date":[str(today)], "Head":[data['head']],"Debit":[None],"Credit":[data['amount']],'Opening Bal':[open],'Closing Bal':[close]}
            else:
                close=open+int(data['amount'])
                new={"Date":[str(today)], "Head":[data['head']],"Debit":[data['amount']*(-1)],"Credit":[None],'Opening Bal':[open],'Closing Bal':[close]}
    elif (data['category']=='Expenses'):
        open=get_close(df[df['Head']==data['head']])
        if data['amount']>0:
            close=open+int(data['amount'])
            new={"Date":[str(today)], "Head":[data['head']],"Debit":[data['amount']],"Credit":[None],'Opening Bal':[open],'Closing Bal':[close]}
        else:
            close=open+int(data['amount'])
            new={"Date":[str(today)], "Head":[data['head']],"Debit":[None],"Credit":[data['amount']*(-1)],'Opening Bal':[open],'Closing Bal':[close]}
    elif(data['category']=='Liabilities'):
        open=get_close(df[df['Head']==data['head']])
        if data['amount']>0:
            close=open+int(data['amount'])
            new={"Date":[str(today)], "Head":[data['head']],"Debit":[None],"Credit":[data['amount']],'Opening Bal':[open],'Closing Bal':[close]}
        else:
            close=open+int(data['amount'])
            new={"Date":[str(today)], "Head":[data['head']],"Debit":[data['amount']*(-1)],"Credit":[None],'Opening Bal':[open],'Closing Bal':[close]}
    else:
        open=get_close(df[df['Head']==data['head']])
        if data['amount']>0:
            close=open+int(data['amount'])
            new={"Date":[str(today)], "Head":[data['head']],"Debit":[None],"Credit":[data['amount']],'Opening Bal':[open],'Closing Bal':[close]}
        else:
            close=open+int(data['amount'])
            new={"Date":[str(today)], "Head":[data['head']],"Debit":[data['amount']*(-1)],"Credit":[None],'Opening Bal':[open],'Closing Bal':[close]}

    df=pd.concat([df,pd.DataFrame(new)],axis=0)
    df.to_csv(filename, index=False)

    df['Date'] = pd.to_datetime(df['Date'])
    df['Year'] = df['Date'].dt.year
    df['Month'] = df['Date'].dt.month
    df['Date'] = df['Date'].dt.date
    months = list(df['Month'].unique())
    if today.month >= 7:
        year1 = today.year
        year2 = year1 + 1
    else:
        year2 = today.year
        year1 = year2 - 1
    if os.path.exists(f'reports/general_voucher_{year1}_{year2}.xlsx'):
        gv=df = pd.read_excel(f'reports/general_voucher_{year1}_{year2}.xlsx', sheet_name=f'{today.month}-{today.year}', header=1, skiprows=[0])
    else:
        gv=pd.DataFrame(columns=['Description','Debit','Credit'])

    # Rename 'Head' column to 'Description'
    df = df.rename(columns={'Head': 'Description'})

    with pd.ExcelWriter(f'reports/general_voucher_{year1}_{year2}.xlsx', engine='xlsxwriter') as writer:
        d = df.iloc[[-1]]
        sheet_name = f'{today.month}-{today.year}'
        d=pd.concat([gv,d[['Description', 'Debit', 'Credit']]],axis=0)

        # Write data to Excel, excluding the 'Date' column
        d.to_excel(writer, sheet_name=sheet_name, index=False, startrow=1)
        
        # Get the xlsxwriter workbook and worksheet objects
        workbook = writer.book
        worksheet = writer.sheets[sheet_name]
        
        # Add today's date to the top left corner
        date_format = workbook.add_format({'num_format': 'yyyy-mm-dd'})
        worksheet.write('A1', f"Date: {today}", date_format)
        
        # Calculate and write totals
        last_row = len(d) + 1
        worksheet.write(last_row + 1, 0, 'Total')
        worksheet.write(last_row + 1, 1, f'=SUM(B2:B{last_row + 1})')
        worksheet.write(last_row + 1, 2, f'=SUM(C2:C{last_row + 1})')

    connection.commit()
    if 'connection' in locals() and connection.is_connected():
        cursor.close()
        connection.close()
    return jsonify({"result": "General Voucher Created Successfully!!!"})

def backup_mysql_database(host, user, password, database, output_file=None, mysqldump_path=None):
    # If no output file is specified, create one with a timestamp
    if output_file is None:
        timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        output_file = f"backup/{database}_backup_{timestamp}.sql"
    try:
        # Print current PATH
        print("Current PATH:")
        for path in os.environ['PATH'].split(os.pathsep):
            print(path)

        # Ensure the directory exists
        os.makedirs(os.path.dirname(output_file), exist_ok=True)

        # Determine mysqldump command
        if mysqldump_path:
            mysqldump_cmd = [mysqldump_path]
        else:
            mysqldump_cmd = ['mysqldump']

        # Construct the full mysqldump command
        mysqldump_cmd.extend([
            f'--host={host}',
            f'--user={user}',
            f'--password={password}',
            '--databases',
            database,
            '--routines',
            '--events',
            '--triggers',
            '--add-drop-table',
            '--create-options'
        ])

        # Print the command being executed (excluding password)
        print("Executing command:", ' '.join(mysqldump_cmd[:2] + mysqldump_cmd[3:]))

        # Execute mysqldump and redirect output to the file
        with open(output_file, 'w') as f:
            subprocess.run(mysqldump_cmd, stdout=f, stderr=subprocess.PIPE, check=True)

        print(f"Database backup created successfully: {output_file}")
    except subprocess.CalledProcessError as e:
        print(f"An error occurred while running mysqldump: {e}")
        print("Error output:", e.stderr.decode())
    except FileNotFoundError:
        print("Error: mysqldump command not found. Please provide the full path to mysqldump using the mysqldump_path parameter.")
        print("Example usage:")
        print("backup_mysql_database(mysqldump_path='C:\\Program Files\\MySQL\\MySQL Server 8.0\\bin\\mysqldump.exe')")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")

@app.route('/generate_records/api', methods=['POST'])
def records():
    connection=create_connection()
    cursor=connection.cursor()

    cursor.execute("SELECT * FROM students;")
    students=cursor.fetchall()
    student_df=pd.DataFrame({"Name":[str(i[0]) for i in students], "FATHER_NAME":[str(i[1]) for i in students], 
                "Gr_no":[str(i[2]) for i in students],"Registration_No":[str(i[3]) for i in students],
                "DOB":[str(i[4]) for i in students],"Contact":[str(i[5]) for i in students],
                "Alternate_Contact":[str(i[6]) for i in students],"Admission_Fees":[str(i[7]) for i in students],
                "Annual_Fund":[i[8] for i in students],"B_Form":[str(i[9]) for i in students],
                "Month_Fees":[str(i[10]) for i in students],"Class":[str(i[11]) for i in students]})

    cursor.execute("SELECT f.ID, f.Registration_No, s.Name, f.Due_date, f.Dues, f.Status, f.advance_months, f.advance_payment, f.month, f.amount_paid, f.advance_payments_last, f.advance_specific_month FROM fee f INNER JOIN students s ON f.Registration_No=s.Registration_No;")
    students=cursor.fetchall()
    student_fee_df=pd.DataFrame({"Payment_ID":[str(i[0]) for i in students], "Registration Number":[str(i[1]) for i in students], 
                "Name":[str(i[2]) for i in students],"Due Date":[str(i[3]) for i in students],
                "Dues":[i[4] for i in students],"Status":[str(i[5]) for i in students],
                "Advance Months":[i[6] for i in students],"Advance Payment":[i[7] for i in students],
                "Month":[i[8] for i in students],"Amount Paid":[i[9] for i in students],
                "Previous Advanced Payment":[i[10] for i in students],"Advance for Specific Month":[i[11] for i in students]})

    cursor.execute("SELECT p.CNIC, t.Name, p.Payment_ID, p.Attendence, p.Total_Payment, p.Medical_Allowance, p.basic_sal, p.deductions, p.advance, p.att_allowance, p.period_engage, p.add_allowance, p.monthh FROM payment p INNER JOIN teachers t ON p.CNIC=t.CNIC;")
    teachers=cursor.fetchall()
    teachers_payment_df=pd.DataFrame({"CNIC":[str(i[0]) for i in students], "Name":[str(i[1]) for i in students], 
                "Payment ID":[str(i[2]) for i in students],"Attendance":[i[3] for i in students],
                "Total Payment":[i[4] for i in students],"Medical Allowance":[i[5] for i in students],
                "Basic Salary":[i[6] for i in students],"Deductions":[i[7] for i in students],
                "Advance":[i[8] for i in students],"Attendance Allowance":[i[9] for i in students],
                "Period Engage":[str(i[10]) for i in students],"Additional Allowance":[i[11] for i in students],
                "Month":[i[11] for i in students]})

    cursor.execute("SELECT * FROM teachers;")
    teachers=cursor.fetchall()
    teachers_df=pd.DataFrame({"CNIC":[str(i[0]) for i in teachers], "Name":[str(i[1]) for i in teachers], 
                "Salary":[str(i[2]) for i in teachers],"CONTACT":[str(i[3]) for i in teachers],
                "Date_OF_Joining":[str(i[4]) for i in teachers]})
    
    with pd.ExcelWriter('datas.xlsx') as writer:
        student_df.to_excel(writer, sheet_name='Students', index=False)
        teachers_df.to_excel(writer, sheet_name='Teachers', index=False)
        student_fee_df.to_excel(writer, sheet_name='Student_fees', index=False)
        teachers_payment_df.to_excel(writer, sheet_name='Teacher Payments', index=False)

    connection.commit()
    if 'connection' in locals() and connection.is_connected():
        cursor.close()
        connection.close()
    
    backup_mysql_database(
        host="localhost",
        user="admin",
        password="root",
        database="arif_academy",
        output_file='backup/data.sql',
        mysqldump_path='C:\\Program Files\\MySQL\\MySQL Server 8.0\\bin\\mysqldump.exe'
    )

    return jsonify({"result": "Data Backup and Generated Successfully!!!"})

@app.route('/misc_expense2/api', methods=['POST'])
def misc_expense2():
    df = pd.read_csv('static/xlscsv/Chart of Accounts.csv', encoding='latin1', usecols=[0,1,2,3,4])
    categories = {
        'Assets': list(df['Assets'].dropna()),
        'Liabilities': list(df['Liabilities'].dropna()),
        'Equity': list(df['Equity'].dropna()),
        'Revenue': list(df['Revenue'].dropna()),
        'Expenses': list(df['Expenses'].dropna())
    }
    return jsonify(categories)

@app.route('/misc_expense_post/api', methods=['POST'])
def misc_expense_post():
    data = request.get_json()
    data=dict(data)
    
    today = datetime.date.today()
    if int(today.month)>6:
        filename=f"static/xlscsv/accounts_{today.year}_{int(today.year)+1}.csv"
    else:
        filename=f"static/xlscsv/accounts_{today.year-1}_{int(today.year)}.csv"
    df=pd.read_csv(filename)
    if data['paymentType']!='Cash':
        data['paymentType']='1010: Current (Bank Account)'
    if (data['category']=='Assets'):
        if 'Contra' in data['head']:
            open=get_close(df[df['Head']==data['paymentType']])
            open2=get_close(df[df['Head']==data['head']])
            close=open+int(data['amount'])
            close2=open2-int(data['amount'])
            new={"Date":[str(today),str(today)], "Head":[data['paymentType'],data['head']],"Debit":[data['amount'], None],"Credit":[None,data['amount']],'Opening Bal':[open,open2],'Closing Bal':[close,close2]}
        else:
            open=get_close(df[df['Head']==data['head']])
            open2=get_close(df[df['Head']==data['paymentType']])
            close=open+int(data['amount'])
            close2=open2-int(data['amount'])
            new={"Date":[str(today),str(today)], "Head":[data['head'],data['paymentType']],"Debit":[data['amount'], None],"Credit":[None,data['amount']],'Opening Bal':[open,open2],'Closing Bal':[close,close2]}
    elif(data['category']=='Equity'):
        if 'Contra' in data['head']:
            open=get_close(df[df['Head']==data['head']])
            open2=get_close(df[df['Head']==data['paymentType']])
            close=open-int(data['amount'])
            close2=open2-int(data['amount'])
            new={"Date":[str(today),str(today)], "Head":[data['head'],data['paymentType']],"Debit":[data['amount'],None],"Credit":[None,data['amount']],'Opening Bal':[open,open2],'Closing Bal':[close,close2]}
        else:
            open=get_close(df[df['Head']==data['paymentType']])
            open2=get_close(df[df['Head']==data['head']])
            close=open+int(data['amount'])
            close2=open2+int(data['amount'])
            new={"Date":[str(today),str(today)], "Head":[data['paymentType'],data['head']],"Debit":[data['amount'],None],"Credit":[None,data['amount']],'Opening Bal':[open,open2],'Closing Bal':[close,close2]}
    elif (data['category']=='Expenses'):
        open=get_close(df[df['Head']==data['head']])
        open2=get_close(df[df['Head']==data['paymentType']])
        close=open+int(data['amount'])
        close2=open2-int(data['amount'])
        new={"Date":[str(today),str(today)], "Head":[data['head'],data['paymentType']],"Debit":[data['amount'], None],"Credit":[None,data['amount']],'Opening Bal':[open,open2],'Closing Bal':[close,close2]}
    elif(data['category']=='Liabilities'):
        open=get_close(df[df['Head']==data['paymentType']])
        open2=get_close(df[df['Head']==data['head']])
        close=open+int(data['amount'])
        close2=open2+int(data['amount'])
        new={"Date":[str(today),str(today)], "Head":[data['paymentType'],data['head']],"Debit":[data['amount'],None],"Credit":[None,data['amount']],'Opening Bal':[open,open2],'Closing Bal':[close,close2]}
    else:
        open=get_close(df[df['Head']==data['paymentType']])
        open2=get_close(df[df['Head']==data['head']])
        close=open+int(data['amount'])
        close2=open2+int(data['amount'])
        new={"Date":[str(today),str(today)], "Head":[data['paymentType'],data['head']],"Debit":[data['amount'],None],"Credit":[None,data['amount']],'Opening Bal':[open,open2],'Closing Bal':[close,close2]}

    df=pd.concat([df,pd.DataFrame(new)],axis=0)
    df.to_csv(filename, index=False)
    return jsonify({"result": "Payment Added Successfully"})

def create_balance_sheet(df, year1,year2,sheet,total):
    assets=df[df['code']==1].sort_values('code')
    liabcap=df[df['code'].isin([2,3])].sort_values('code')
    assets.fillna(0,inplace=True)
    assets['Debit']=assets['Debit']-assets['Credit']
    assets=assets.drop(columns=['Credit','diff','code'],axis=1).reset_index(drop=True)
    liabcap.fillna(0,inplace=True)
    liabcap['Credit']=liabcap['Credit']-liabcap['Debit']
    liabcap.drop(columns=['Debit','diff','code'],axis=1,inplace=True)
    liabcap.rename(columns={'Head':'Head2'},inplace=True)
    liabcap=pd.concat([liabcap,pd.DataFrame({'Head2':['Profit & Loss'],'Credit':[total]})],axis=0).reset_index(drop=True)
    balance_sheet=pd.concat([assets,liabcap],axis=1)
    balance_sheet=pd.concat([balance_sheet,pd.DataFrame({'Debit':[balance_sheet['Debit'].sum()],'Credit':[balance_sheet['Credit'].sum()]})])
    balance_sheet['Debit']=balance_sheet['Debit'].astype('str')
    balance_sheet['Credit']=balance_sheet['Credit'].astype('str')
    balance_sheet['Debit']=np.where((balance_sheet['Debit'].str[0]=='-'),'('+balance_sheet['Debit'].str[1:]+')',balance_sheet['Debit'])
    balance_sheet['Credit']=np.where((balance_sheet['Credit'].str[0]=='-'),'('+balance_sheet['Credit'].str[1:]+')',balance_sheet['Credit'])
    balance_sheet=balance_sheet.apply(lambda x: x.str.replace('nan',''))
    return balance_sheet

def generate_cashbook(df):
    df=df[['Date', 'Head','Debit','Credit','Opening Bal','Closing Bal']]
    cash=df[df['Head']=='Cash']
    bank=df[df['Head']=='1010: Current (Bank Account)']
    if len(cash)>0:
        cash=pd.concat([cash,pd.DataFrame({'Credit':['Opening Balance', 'Closing Balance'],'Closing Bal':[int(cash.iloc[0,4]),int(cash.iloc[-1,5])]})])
    if len(bank)>0:
        bank=pd.concat([bank,pd.DataFrame({'Credit':['Opening Balance', 'Closing Balance'],'Closing Bal':[int(bank.iloc[0,4]),int(bank.iloc[-1,5])]})])
    cc_bb=[cash,bank]
    return cc_bb

def create_trial_balance(df):
    df['code']=df['Head'].str[0]
    df['code']=np.where(df['code']=='C','1',df['code'])
    df['code']=df['code'].astype('int64')
    cash_book=generate_cashbook(df)
    df=df.groupby(['Head','code'])[['Debit','Credit']].sum().reset_index()
    df['diff']=df['Debit']-df['Credit']
    df['Debit']=np.where(df['diff']>0,df['diff'],np.nan)
    df['Credit']=np.where(df['diff']<0,df['diff']*(-1),np.nan)
    df=df.sort_values('Debit')
    trial_bal=df[['Head','Debit','Credit']]
    trial_bal=pd.concat([trial_bal,pd.DataFrame({'Head':['Total'],'Debit':[df['Debit'].sum()],'Credit':[df['Credit'].sum()]})],axis=0)
    return df, trial_bal, cash_book

def create_income_statement(df, year1,year2,sheet,month=None):
    if month is not None:
        d=df[df['Month'].isin(month)]
    else:
        d=df
    df, trial_bal, cash_book=create_trial_balance(d)
    income_statement=df[(df['code']==4)|(df['code']>=5)].sort_values('code')
    income_statement.drop(columns=['diff'],axis=1,inplace=True)
    income_statement=pd.concat([income_statement,pd.DataFrame({'Head':['Total Exppenditure','Net Income / Loss'],'Credit':[-1*income_statement['Debit'].sum(),income_statement['Credit'].sum()-income_statement['Debit'].sum()]})],axis=0)
    income_statement.fillna(0,inplace=True)
    total=int(list(income_statement.loc[income_statement['Head']=='Net Income / Loss','Credit'])[0])
    income_statement['Debit']=income_statement['Debit'].astype('str')
    income_statement['Credit']=income_statement['Credit'].astype('str')
    income_statement['Debit']=np.where(((income_statement['Debit'].str[0]=='-')|(income_statement['Debit']=='0.0')),income_statement['Debit'].str[1:],'('+income_statement['Debit'].str[:]+')')
    income_statement['Credit']=np.where(((income_statement['Credit'].str[0]=='-')&(income_statement['Credit']!='0.0')),'('+income_statement['Credit'].str[1:]+')',income_statement['Credit'])
    income_statement['Debit']=np.where(income_statement['Debit'].isin(['.0','0.0']),"",income_statement['Debit'])
    income_statement['Credit']=np.where(income_statement['Credit'].isin(['.0','0.0']),"",income_statement['Credit'])
    income_statement.drop(columns=['code'],axis=1,inplace=True)
    bal_sheet=create_balance_sheet(df,year1,year2,sheet,total)
    return (trial_bal,income_statement,bal_sheet,cash_book,sheet)

def generate_cashflows(df2,df1, year1, year2):
    def convert_to_float(val):
        if isinstance(val, str):
            if val!='':
                if val.startswith('(') and val.endswith(')'):
                    return -float(val.strip('()').replace(',', ''))
                else:
                    return float(val.replace(',', ''))
        return val
    # Convert string columns to float
    for df in [df1, df2]:
        for col in ['Credit', 'Debit']:
            if df[col].dtype == 'object':
                df[col] = df[col].apply(convert_to_float)
    
    df1=df1.iloc[:-1]
    df2=df2.iloc[:-1]
    df1=pd.concat([df1[['Head','Debit']],df1[['Head2','Credit']].rename(columns={'Head2':'Head','Credit':'Debit'})],axis=0)
    df2=pd.concat([df2[['Head','Debit']],df2[['Head2','Credit']].rename(columns={'Head2':'Head','Credit':'Debit'})],axis=0)
    df1.dropna(inplace=True)
    df2.dropna(inplace=True)

    # Combine DataFrames
    combined_df = pd.concat([df1.set_index('Head'), df2.set_index('Head')], axis=1, keys=['Current', 'Previous'])
    combined_df = combined_df.fillna(0)


    # Calculate changes
    changes = combined_df['Current'] - combined_df['Previous']

    # Define account classifications
    op = [
        '1210: Accounts Receivable', '1310: Inventory', '2010: Accounts Payable',
        '2115: Accrued Employee Benefits', '1220: Fee Receivable', '1320: Stationery Unused/ Stock',
        '1330: Printed Supplies Unused/ Stock', '1410: Prepaid Expenses', '2100: Accrued Salaries',
        '2110: Accrued daily Wages', '2150: Accrued Any taxes and duties', '2200: Advanced Fee Received',
        '2250: Rent Payable', '2260: Utilities Payable', '2300: Misc. Services Payable'
    ]
    
    iv = [
        '1510: Property / Building', '1520: Furniture & Fixtures', '1530: Lab. Equipment (Electronic)/ Mechanical',
        '1540: Equipment (Electronic)/ Mechanical', '1590: Accumulated Depreciation (Contra Assets)',
        '1690: Accumulated Amortization (Contra Assets)'
    ]
    
    fa = [
        '3100: Capital', '3200: Profit or Loss Reserve A/c', '3300: Drawings (Contra Equity)'
    ]

    operating_accounts = [i for i in op if i in list(combined_df.index)]
    investing_accounts = [i for i in iv if i in list(combined_df.index)]
    financing_accounts = [i for i in fa if i in list(combined_df.index)]

    print(operating_accounts)

    # Calculate cash flows
    operating_cash_flow = -changes.loc[operating_accounts].sum()
    investing_cash_flow = -changes.loc[investing_accounts].sum()
    financing_cash_flow = changes.loc[financing_accounts].sum()

    # Calculate net change in cash
    net_change_in_cash = operating_cash_flow + investing_cash_flow + financing_cash_flow

    # Create cash flow statement DataFrame
    cash_flow_statement = pd.DataFrame({
        'Activity': ['Operating Activities', 'Investing Activities', 'Financing Activities', 'Net Change in Cash'],
        'Amount': [int(operating_cash_flow), int(investing_cash_flow), int(financing_cash_flow), int(net_change_in_cash)]
    })

    return cash_flow_statement

@app.route('/generate_report/api', methods=['POST'])
def report():
    connection=create_connection()
    cursor=connection.cursor()

    data=request.get_json()
    data=dict(data)

    d_m=data['d_m']
    val=data['val']

    today = datetime.date.today()
    if int(today.month)>6:
        filename=f"static/xlscsv/accounts_{today.year}_{int(today.year)+1}.csv"
    else:
        filename=f"static/xlscsv/accounts_{today.year-1}_{int(today.year)}.csv"
    df=pd.read_csv(filename)
    
    df['Date']=pd.to_datetime(df['Date'])

    if d_m == 'days':
        start_date = today - datetime.timedelta(days=val)
    elif d_m == 'months':
        start_date = today - pd.DateOffset(months=val)
    else:
        raise ValueError("Invalid value for d_m. Must be 'days' or 'months'.")
    
    df['Month']=df['Date'].dt.month
    df['Year']=df['Date'].dt.year
    if today.month>=7:
        year1=today.year
        year2=year1+1
    else:
        year2=today.year
        year1=year2-1
    trial_bals,income_statements,bal_sheets,cash_books=[],[],[],[]
    df['today']=today
    df['today']=pd.to_datetime(df['today'])
    df['start_date']=start_date
    df['start_date']=pd.to_datetime(df['start_date'])
    filtered_df = df[(df['Date'] >= df['start_date']) & (df['Date'] <= df['today'])]
    
    trial_bal,income_statement,bal_sheet,cash_book,sheet=create_income_statement(filtered_df, year1,year2,f'{str(today)}_{val}_{d_m}')
    trial_bals.append((trial_bal,sheet))
    income_statements.append((income_statement,sheet))
    bal_sheets.append((bal_sheet,sheet))
    cash_books.append((cash_book,sheet))

    for month in list(df['Month'].unique()):
        trial_bal,income_statement,bal_sheet,cash_book,sheet=create_income_statement(df, year1,year2,f'{month}_{list(df['Year'].unique())[0]}', month=[month])
        trial_bals.append((trial_bal,sheet))
        income_statements.append((income_statement,sheet))
        bal_sheets.append((bal_sheet,sheet))
        cash_books.append((cash_book,sheet))

    quarter=[[7,8,9],[10,11,12],[1,2,3],[4,5,6]]
    for i,q in enumerate(quarter):
        exist=True
        for month in q:
            if month not in list(df['Month'].unique()):
                exist=False
        if exist==False:
            break
        trial_bal,income_statement,bal_sheet,cash_book,sheet=create_income_statement(df,  year1,year2,f'q{i}',month=q)
        trial_bals.append((trial_bal,sheet))
        income_statements.append((income_statement,sheet))
        bal_sheets.append((bal_sheet,sheet))
        cash_books.append((cash_book,sheet))
        
    semi=[[7,8,9,10,11,12],[1,2,3,4,5,6]]
    for i,s in enumerate(semi):
        exist=True
        for month in q:
            if month not in list(df['Month'].unique()):
                exist=False
        if exist==False:
            break
        trial_bal,income_statement,bal_sheet,cash_book,sheet=create_income_statement(df,  year1,year2,f's{i}',month=s)
        trial_bals.append((trial_bal,sheet))
        income_statements.append((income_statement,sheet))
        bal_sheets.append((bal_sheet,sheet))
        cash_books.append((cash_book,sheet))
    trial_bal,income_statement,bal_sheet,cash_book,sheet=create_income_statement(df, year1,year2,f'yearly',month=list(df['Month'].unique()))
    file_path = f'reports/Balance_Sheet_{year1-1}_{year1}.xlsx'
    if os.path.exists(file_path):
        bal1=pd.read_excel(file_path,sheet_name='yearly')
        cashflows=generate_cashflows(bal1,bal_sheet, year1, year2)
        with pd.ExcelWriter(f'reports/Cash_Flows_{year1}_{year2}.xlsx') as writer:
            cashflows.to_excel(writer, sheet_name='cash', index=False)
    trial_bals.append((trial_bal,sheet))
    income_statements.append((income_statement,sheet))
    bal_sheets.append((bal_sheet,sheet))
    cash_books.append((cash_book,sheet))

    with pd.ExcelWriter(f'reports/Trial_Balance_{year1}_{year2}.xlsx') as writer:
        for i in trial_bals:
            i[0].to_excel(writer, sheet_name=i[1], index=False)
    with pd.ExcelWriter(f'reports/Income_Statement_{year1}_{year2}.xlsx') as writer:
        for i in income_statements:
            i[0].to_excel(writer, sheet_name=i[1], index=False)
    with pd.ExcelWriter(f'reports/Balance_Sheet_{year1}_{year2}.xlsx') as writer:
        for i in bal_sheets:
            i[0].to_excel(writer, sheet_name=i[1], index=False)
    with pd.ExcelWriter(f'reports/Cash_Book_{year1}_{year2}.xlsx') as writer:
        for i in cash_books:
            i[0][0].to_excel(writer, sheet_name=f"cash_{i[1]}", index=False)
            i[0][1].to_excel(writer, sheet_name=f"bank_{i[1]}", index=False)

    connection.commit()
    if 'connection' in locals() and connection.is_connected():
        cursor.close()
        connection.close()
    return jsonify({"result": "Reports Generated Successfully!!"})

def restore_mysql_database(host, user, password, database, input_file, mysql_path=None):
    try:
        # Connect to the MySQL server
        conn = mysql.connector.connect(
            host=host,
            user=user,
            password=password,
            database=database
        )
        cursor = conn.cursor()

        # Disable foreign key checks
        cursor.execute("SET FOREIGN_KEY_CHECKS = 0;")

        # Get all tables in the database
        cursor.execute("SHOW TABLES;")
        tables = [table[0] for table in cursor.fetchall()]

        # Drop all tables
        for table in tables:
            cursor.execute(f"DROP TABLE IF EXISTS `{table}`;")

        # Enable foreign key checks
        cursor.execute("SET FOREIGN_KEY_CHECKS = 1;")

        conn.commit()
        cursor.close()
        conn.close()

        # Determine mysql command
        if mysql_path:
            mysql_cmd = [mysql_path]
        else:
            mysql_cmd = ['mysql']

        # Construct the full mysql command
        mysql_cmd.extend([
            f'--host={host}',
            f'--user={user}',
            f'--password={password}',
            database
        ])

        # Check if input file exists
        if not os.path.exists(input_file):
            raise FileNotFoundError(f"Input file not found: {input_file}")

        # Execute the SQL file to restore the database
        with open(input_file, 'r') as f:
            subprocess.run(mysql_cmd, stdin=f, check=True)

        print("Database restored successfully.")
    except FileNotFoundError as e:
        print(f"File not found error: {e}")
        print("If mysql command is not found, please provide the full path to mysql using the mysql_path parameter.")
        print("Example usage:")
        print("restore_mysql_database(mysql_path='C:\\Program Files\\MySQL\\MySQL Server 8.0\\bin\\mysql.exe')")
    except subprocess.CalledProcessError as e:
        print(f"An error occurred while running mysql: {e}")
        print("Error output:", e.stderr.decode() if e.stderr else "No error output available")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")


@app.route('/upload_xls/api', methods=['POST'])
def uploadd():
    restore_mysql_database(
        host="localhost",
        user="admin",
        password="root",
        database="arif_academy",
        input_file='backup/data.sql',
        mysql_path='C:\\Program Files\\MySQL\\MySQL Server 8.0\\bin\\mysql.exe'
    )
    return jsonify({"result": "Database Updated Successfully!!!"})

@app.route('/chart/api', methods=['POST'])
def chartt():
    data = request.get_json()
    data=dict(data)

    head = data['head']
    category=data['category']
    account=data['account']
    action=data['action']
    df=pd.read_csv('static/xlscsv/Chart of Accounts.csv',usecols=[0,1,2,3,4])
    if (action=='add'):
        df=pd.concat([df, pd.DataFrame({category:[np.nan]})],axis=0).reset_index(drop=True)
        df[category]=np.where(df.index==(len(df)-1), account,df[category])
    else:
        df[category]=np.where(df[category]==head,np.nan,df[category])
    df.to_csv('static/xlscsv/Chart of Accounts.csv',index=False)
    return jsonify({"result": "Charts updated Successfully!!!"})

def open_browser():
    webbrowser.open_new('http://127.0.0.1:5000/')

if __name__ == '__main__':
    Timer(1, open_browser).start()
    app.run(debug=True)