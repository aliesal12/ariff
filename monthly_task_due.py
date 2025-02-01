import os
import datetime
import time
import subprocess
from flask import Flask, render_template, request, jsonify
from os import getenv
import numpy as np
from datetime import datetime
import mysql.connector
from mysql.connector import Error
import pandas as pd
import datetime as dt
from docx import Document
from docx.shared import Pt
import math

def create_connection():
    try:
        connection = mysql.connector.connect(
            host='localhost',
            user='admin',
            password='root',
            database='arif_academy'
        )
        
        if connection.is_connected():
            return connection
                
    except Error as e:
        print("Error while connecting to MySQL", e)

"""def should_run():
    today = datetime.date.today()
    state_file = 'last_run_state.txt'
    
    if not os.path.exists(state_file):
        return True
    
    with open(state_file, 'r') as f:
        last_run = datetime.datetime.strptime(f.read().strip(), '%Y-%m-%d').date()
    
    return today.month != last_run.month or today.day == 1"""

def update_state():
    with open('last_run_state.txt', 'w') as f:
        f.write(dt.date.today().strftime('%Y-%m-%d'))

def generate_notice(student_id, name, classs, dues, monthly_fee,today):
    print(student_id)
    # Create a new Document
    doc = Document()

    # Add date
    current_date = datetime.now().strftime("%d-%m-%Y")
    doc.add_paragraph(f"Date: {current_date}")
    doc.add_paragraph("")
    if int(dues)!=0:
        months=int(math.ceil(int(dues)/int(monthly_fee)))
    else:
        months=1
    month={1:"January", 2:'February',3:'March',4:'April',5:'May',6:'June',7:'July',8:'August',9:'September',10:'October',11:'November',12:'December'}
    if months==1:
        add_month=month[int(today.month)]
    elif months==2:
        if int(today.month)-1==0:
            add_month=f'{"December"} and {month[int(today.month)]}'
        else:
            add_month=f'{month[int(today.month)-1]} and {month[int(today.month)]}'
    else:
        if int(today.month)-1==0:
            add_month=f'{"December"} and {month[int(today.month)]}'
        else:
            add_month=f'{month[int(today.month)-1]} and {month[int(today.month)]}'
            for i in range(3,months+1):
                monthh=int(today.month)-i
                if monthh<=0:
                    monthh=monthh+12
                add_month=f'{month[monthh]}, {add_month}'
    # Add main content
    doc.add_paragraph("Dear Parents,")
    doc.add_paragraph("")
    doc.add_paragraph(
        f"This is to inform you that the Tuition fee, admission fee (For new admission) and "
        f"Annual Fund of your child named {name} currently in class {classs} "
        f"are overdue for the months of {add_month} amounting to Pkr {dues} "
        f"(in total) @ monthly fee of Pkr {monthly_fee}, if the dues are not cleared by {25}-{today.month}-{today.year} "
        f"date, the name of your child shall be struck off and re-admission policy shall be applied subject to the "
        f"clearance of the above dues."
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "This is the final notice to you, kindly contact school administration and clear the above dues immediately "
        "to avoid the restriction of your child / ward from AAE."
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "Note: If you find any discrepancy in this respect, please bring the school card to update the record before the last date of payment mentioned above."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Administration - AAE")
    

    if not os.path.exists('student_data'):
        os.makedirs('student_data')
    # Save the filled document
    output_path = f'student_data/output_notice_{student_id.replace('/','')}.docx'
    doc.save(output_path)
    
    print(f"Notice generated and saved as {output_path}")
    return True
def main_program():
    print("Running the main program...")
    connection=create_connection()
    cursor=connection.cursor()
    today=dt.date.today()
    month=int(today.month)
    cursor.execute(f'SELECT s.Registration_no, s.Name, s.Month_Fees, f.Dues, f.advance_payment, f.amount_paid, f.Status, f.advance_payments_last,f.advance_months, s.Class FROM students s LEFT JOIN (SELECT * FROM fee WHERE month={month}) f ON s.Registration_No=f.Registration_No WHERE f.Status="Unpaid";')
    data=cursor.fetchall()
    data_df=pd.DataFrame({"Reg_No":[str(i[0]) for i in data], "Name":[str(i[1]) for i in data], 
                    "Monthly_fee":[str(i[2]) for i in data],"Dues":[str(i[3]) for i in data],
                    "advance_payment":[str(i[4]) for i in data],"amount_paid":[str(i[5]) for i in data],
                    'status':[str(i[6]) for i in data],'advance_payments_last':[str(i[7]) for i in data],
                    'advance_months':[str(i[8]) for i in data],'Class':[str(i[9]) for i in data]})
    data_df['advance_payment']=np.where((data_df['advance_payment']=='None')|(data_df['advance_payment'].isna()), 0, data_df['advance_payment'])
    data_df['advance_payment']=data_df['advance_payment'].astype('int64')
    data_df['amount_paid']=np.where((data_df['amount_paid']=='None')|(data_df['amount_paid'].isna()), 0, data_df['amount_paid'])
    data_df['amount_paid']=data_df['amount_paid'].astype('int64')
    data_df['Dues']=np.where((data_df['Dues']=='None')|(data_df['Dues'].isna()), 0, data_df['Dues'])
    data_df['Dues']=data_df['Dues'].astype('int64')
    data_df['advance_payments_last']=np.where((data_df['advance_payments_last']=='None')|(data_df['advance_payments_last'].isna()), 0, data_df['advance_payments_last'])
    data_df['advance_payments_last']=data_df['advance_payments_last'].astype('int64')
    data_df['advance_months']=np.where((data_df['advance_months']=='None')|(data_df['advance_months'].isna()), 0, data_df['advance_months'])
    data_df['advance_months']=data_df['advance_months'].astype('int64')
    data_df['Monthly_fee']=data_df['Monthly_fee'].astype('int64')
    data_df['Monthly_fee']=data_df['Monthly_fee'].fillna(0)
    data_df=data_df[(data_df['Dues'].notnull())&(data_df['Dues']>0)]
    data_df=data_df.fillna(0)
    if not os.path.exists('student_data'):
        os.makedirs('student_data')
    data_df.to_csv('student_data/students_dues.csv')
    print(data_df.columns)
    data_df['generated']=data_df.apply(lambda row: generate_notice(row['Reg_No'], row['Name'], row['Class'], row['Dues'], row['Monthly_fee'],today),axis=1)

    time.sleep(10)  # Wait for 10 seconds
    print("Main program completed.")

def run_if_needed():
    #if should_run():
    try:
        main_program()
        update_state()
    except Exception as e:
        print(f"An error occurred: {e}")
            # Don't update the state file, so it will run again on next startup

if __name__ == "__main__":
    run_if_needed()